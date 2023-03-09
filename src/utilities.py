import ast
import re
import logging
import os
import docx
from docx.shared import Pt
import pandas as pd

logger = logging.getLogger(__name__)


def is_question(text: str):
    return text[0].isdigit()


def convert_question_columns_to_number(df: pd.DataFrame):
    df.columns = [get_number_from_question(column) for column in df.columns]
    return df


def get_question_with_number(number, question_list):
    for question in question_list:
        if get_number_from_question(question) == number:
            return question


def intersection(lst1, lst2):
    return set(lst1).intersection(lst2)


def get_contradiction(list_of_numbers, cols, total_contradict_score):
    from itertools import permutations

    pairs = permutations(list_of_numbers, 2)
    for pair in pairs:
        contradict_score = abs(int(pair[0]) - int(pair[1]))
        if contradict_score > 0:
            total_contradict_score += contradict_score
            return {'cols': set(cols), 'contradiction_score': total_contradict_score}
    return {'cols': [], 'contradiction_score': total_contradict_score}


def get_number_from_question(text: str):
    pattern = re.compile('^\s*[0-9]+')
    if text[0].isdigit():
        try:
            return int(pattern.search(text).group(0))
        except:
            logger.error(f"Can not parse question: {text}")
    else:
        return text


def calculate_avg(list_of_integers):
    sum = 0
    for i in list_of_integers:
        sum += i

    return sum / len(list_of_integers)


def mean_highlighter(x):
    style_lt = "background-color: #EE2E31; color: white; font-weight: bold;"
    style_gt = "background-color: #31D843; color: white; font-weight: bold;"
    gt_mean = x > x.mean()
    return [str(style_gt) if i else str(style_lt) for i in gt_mean]


def get_question_text_with_num(num_list, question_list):
    return_dict = {}
    for num in num_list:
        for question in question_list:
            if get_number_from_question(question) == num:
                return_dict[num] = question
    return return_dict


def check_contradiction(list_of_numbers):
    from itertools import permutations

    pairs = permutations(list_of_numbers, 2)
    for pair in pairs:
        if abs(int(pair[0]) - int(pair[1])) > 0:
            return True
    return False


def get_contradict_list(num_to_check, list_of_contradict_list):
    for list in list_of_contradict_list:
        if num_to_check in list:
            return list
    return None


def flat_text_map(text_map):
    return_dict = {}
    for k, v in text_map.items():
        for k2, v2 in v.items():
            return_dict[f"{k}_{k2}"] = v2
    return return_dict


def create_text_report(input_dir, output_dir):
    import os
    dir_path = input_dir
    files = []
    for file in os.listdir(dir_path):
        if file.endswith("text.xlsx"):
            files.append(file)
    ordered_list = ['b5kt-text.xlsx', 'cmvkb-text.xlsx', 'cipto-text.xlsx']
    paths = []
    for file in ordered_list:
        if file in files:
            paths.append(f"{dir_path}/{file}")

    contradict_files = []
    for file in os.listdir(dir_path):
        if file.endswith("contradict.xlsx"):
            contradict_files.append(file)
    contradict_paths = [f"{dir_path}/{file}" for file in contradict_files]

    all_df = pd.read_excel(
        f"{dir_path}/all_reports.xlsx",
        engine='openpyxl')
    all_df.index = all_df['İsim Soyisim']
    all_df = all_df.loc[~all_df.index.duplicated(keep='first')]
    all_df.drop(['İsim Soyisim'], axis=1, inplace=True)

    for path in paths:
        df = pd.read_excel(path, engine='openpyxl')
        df.index = df['İsim Soyisim']
        df.drop(['İsim Soyisim'], axis=1, inplace=True)
        df = df.loc[~df.index.duplicated(keep='first')]
        all_df = pd.concat([all_df, df], axis=1)

    for path in contradict_paths:
        df = pd.read_excel(path, engine='openpyxl')
        df.index = df['İsim Soyisim']
        df.drop(['İsim Soyisim'], axis=1, inplace=True)
        df = df.loc[~df.index.duplicated(keep='first')]
        all_df = pd.concat([all_df, df], join="inner", axis=1)

    from docx import Document
    # create the docx report
    columns = all_df.columns
    columns = [column for column in columns if "contradict" not in column]
    for index, row in all_df.iterrows():
        document = Document()
        document.add_heading(index, 0)
        for column in columns:
            document.add_heading(column, level=2)
            document.add_paragraph(str(row[column]))
            if not isNaN(row[column+" contradict"]):
                contradicts = row[column+" contradict"]
                contradicts = ast.literal_eval(contradicts)
                for contradict in contradicts:
                    paragraph = document.add_paragraph('')
                    for cont in contradict.split("Soru:"):
                        sentence = paragraph.add_run(cont)
                        sentence.font.name = 'Arial'
                        sentence.font.size = docx.shared.Pt(8)

        document.save(f"{output_dir}/{index}.docx")


def isNaN(string):
    return string != string


def create_dirs(path):
    if os.path.isdir(f"{path}/raporlar"):
        import shutil

        shutil.rmtree(f"{path}/raporlar")

    os.mkdir(f"{path}/raporlar")
    os.mkdir(f"{path}/raporlar/text_reports")
    os.mkdir(f"{path}/raporlar/excel_reports")
    os.mkdir(f"{path}/raporlar/temp_files")
    os.mkdir(f"{path}/raporlar/text_reports/person_reports")
    os.mkdir(f"{path}/raporlar/excel_reports/mlq_reports")


def common_paths(path):
    return \
        f"{path}/../resources/yanıtlar/", \
        f"{path}/raporlar/excel_reports", \
        f"{path}/raporlar/text_reports", \
        f"{path}/raporlar/temp_files", \
        f"{path}/../resources/passwords.json"
